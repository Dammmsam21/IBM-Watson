#!/usr/bin/env python3
"""Convert the markdown study-guide files into a single formatted Word document."""
import re
import glob
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GUIDE_DIR = "study-guide"
OUT = "Deloitte-AA-Research-Tools-Study-Guide.docx"

# Order of files in the combined document
ORDER = [
    "00-overview-and-how-to-use.md",
    "01-factiva.md",
    "02-onesource.md",
    "03-dnb-hoovers.md",
    "04-thomson-lseg-workspace.md",
    "05-boardex.md",
    "06-sp-capital-iq.md",
    "07-pitchbook.md",
    "08-bloomberg.md",
    "09-ibisworld.md",
    "10-microsoft-office.md",
    "11-core-concepts-and-domain.md",
    "12-glossary.md",
    "13-study-plan-and-interview-prep.md",
]

ACCENT = RGBColor(0x0B, 0x5A, 0x3C)   # deloitte-ish green
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`|\[.+?\]\(.+?\))")


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_inline(paragraph, text):
    """Add text to a paragraph, parsing **bold**, `code`, and [link](url)."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(re.sub(r"\*([^*\n]+?)\*", r"\1", part[2:-2]))
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[(.+?)\]\((.+?)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def strip_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*([^*\n]+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
    return text


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = []
    for r in rows[2:]:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        body.append(cells)
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_inline(p, h)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
        shade_cell(hdr[i], "0B5A3C")
    for cells in body:
        row = table.add_row().cells
        for i in range(ncols):
            val = cells[i] if i < len(cells) else ""
            row[i].text = ""
            p = row[i].paragraphs[0]
            add_inline(p, val)
            for run in p.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "0B5A3C")
    pbdr.append(left)
    pPr.append(pbdr)
    add_inline(p, text)
    for run in p.runs:
        run.italic = True
        run.font.color.rgb = GREY


def process_lines(doc, lines):
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+$", stripped):
            add_hr(doc)
            i += 1
            continue

        # table (line with | and next line is separator)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1].strip()):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_table(doc, tbl)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                h = doc.add_heading(level=1)
                run = h.add_run(strip_inline(text))
                run.font.color.rgb = ACCENT
            elif level == 2:
                h = doc.add_heading(level=2)
                run = h.add_run(strip_inline(text))
                run.font.color.rgb = ACCENT
            else:
                h = doc.add_heading(level=min(level, 4))
                run = h.add_run(strip_inline(text))
                run.font.color.rgb = DARK
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            add_blockquote(doc, quote)
            i += 1
            continue

        # checkbox list
        cb = re.match(r"^- \[ \]\s+(.*)$", stripped)
        if cb:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, cb.group(1))
            i += 1
            continue

        # ordered list
        ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ol:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, ol.group(2))
            i += 1
            continue

        # bullet list (support one level of nesting via indentation)
        bl = re.match(r"^(\s*)[-*]\s+(.*)$", line.rstrip("\n"))
        if bl:
            indent = len(bl.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, bl.group(2))
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Study Guide")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = ACCENT

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Research Tools for the Deloitte A&A\nStrategic Relationship Management Role")
    r.font.size = Pt(16)
    r.font.color.rgb = DARK

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Associate Consultant, Audit & Assurance Growth (USI, Hyderabad)  -  Req. 355084")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Factiva  -  OneSource  -  D&B Hoovers  -  Thomson / LSEG Workspace  -  BoardEx\n"
                     "S&P Capital IQ  -  PitchBook  -  Bloomberg  -  IBISWorld  -  Microsoft Office")
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    doc.add_page_break()


def main():
    doc = Document()
    set_base_styles(doc)
    add_cover(doc)

    # Table of contents (static list)
    h = doc.add_heading(level=1)
    run = h.add_run("Contents")
    run.font.color.rgb = ACCENT
    toc_items = [
        "Overview & How to Use This Guide",
        "1. Dow Jones Factiva",
        "2. OneSource",
        "3. D&B Hoovers",
        "4. Thomson / LSEG Workspace",
        "5. BoardEx",
        "6. S&P Capital IQ",
        "7. PitchBook",
        "8. Bloomberg",
        "9. IBISWorld",
        "10. Microsoft Office (PowerPoint, Excel, Word)",
        "11. Core Concepts & Domain Knowledge",
        "12. Glossary & Name-Change Cheat Sheet",
        "13. Study Plan & Interview Prep",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)
    doc.add_page_break()

    for idx, fname in enumerate(ORDER):
        path = os.path.join(GUIDE_DIR, fname)
        with open(path, encoding="utf-8") as f:
            lines = f.read().splitlines()
        process_lines(doc, lines)
        if idx != len(ORDER) - 1:
            doc.add_page_break()

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
